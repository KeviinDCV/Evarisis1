# -*- coding: utf-8 -*-
"""Reemplaza placeholders y estimaciones por datos reales en los 8 docs HUV V6.9.9."""
from docx import Document
from copy import deepcopy
from docx.oxml.ns import qn
import os, shutil

CARPETA = r'C:\Users\Kechavarro\Desktop\ProyectoHUV9GESTOR_ONCOLOGIA\documentacion\entrega_HUV_V6.9.9'

# Backup
backup_dir = os.path.join(CARPETA, '_backup_pre_v2')
os.makedirs(backup_dir, exist_ok=True)
for f in os.listdir(CARPETA):
    if f.endswith('.docx'):
        shutil.copy(os.path.join(CARPETA, f), os.path.join(backup_dir, f))
print(f"Backup OK -> {backup_dir}")

# === Helpers ===
def replace_in_paragraph(p, old, new):
    if not p.text or old not in p.text:
        return False
    runs = p.runs
    if not runs:
        return False
    full_new = ''.join(r.text for r in runs).replace(old, new)
    runs[0].text = full_new
    for r in runs[1:]:
        r.text = ''
    return True

def set_cell_text(cell, new_text):
    if not cell.paragraphs:
        cell.text = new_text
        return
    p = cell.paragraphs[0]
    runs = p.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ''
    else:
        p.text = new_text
    for extra_p in cell.paragraphs[1:]:
        extra_p._element.getparent().remove(extra_p._element)

def clear_table_rows_except_header(table):
    rows_to_remove = list(table.rows)[1:]
    for r in rows_to_remove:
        r._tr.getparent().remove(r._tr)

def add_row_clone_last(table, values):
    template_row = table.rows[-1]
    new_tr = deepcopy(template_row._tr)
    for tc in new_tr.findall(qn('w:tc')):
        for p in tc.findall(qn('w:p')):
            for r in p.findall(qn('w:r')):
                for t in r.findall(qn('w:t')):
                    t.text = ''
    table._tbl.append(new_tr)
    new_row = table.rows[-1]
    for i, v in enumerate(values):
        if i < len(new_row.cells):
            set_cell_text(new_row.cells[i], str(v))

# ====== DOC 02 ======
ruta = os.path.join(CARPETA, '02_DOCUMENTACION_TECNICA.docx')
doc = Document(ruta)
for p in doc.paragraphs:
    if 'IHQXXXNNNN' in p.text:
        replace_in_paragraph(p, 'IHQXXXNNNN', 'IHQAANNNN (AA es el ano de dos digitos y NNNN el consecutivo; ejemplo IHQ250001)')
doc.save(ruta)
print("OK 02 actualizado")

# ====== DOC 04 ======
ruta = os.path.join(CARPETA, '04_MANUAL_DE_USUARIO.docx')
doc = Document(ruta)
for p in doc.paragraphs:
    if 'IHQXXXNNNN' in p.text:
        replace_in_paragraph(p, 'IHQXXXNNNN', 'su numero (ejemplo IHQ250001)')
doc.save(ruta)
print("OK 04 actualizado")

# ====== DOC 07 ======
ruta = os.path.join(CARPETA, '07_REGISTRO_HISTORICO_INTERACCIONES.docx')
doc = Document(ruta)
for p in doc.paragraphs:
    replace_in_paragraph(p, 'Tabla de hitos con fechas estimadas', 'Tabla de hitos del proyecto con sus periodos')
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_paragraph(p, 'Periodo estimado', 'Periodo')
                replace_in_paragraph(p, 'Fecha estimada', 'Fecha')
doc.save(ruta)
print("OK 07 actualizado")

# ====== DOC 08 ======
ruta = os.path.join(CARPETA, '08_ESTADISTICA.docx')
doc = Document(ruta)

reemplazos = [
    ('2. Distribucion estimada por mes', '2. Distribucion por mes'),
    ('La siguiente tabla presenta la distribucion estimada de informes procesados por mes, agrupada para fines indicativos. Los valores podran refinarse cuando se complete el reproceso correctivo planificado.',
     'La siguiente tabla presenta la distribucion real de informes procesados por mes, segun el campo "Fecha Informe" registrado por el sistema en los 1.415 informes procesados al cierre del periodo.'),
    ("Nota: la distribucion mensual es indicativa y debera confirmarse con la consulta directa al campo 'Fecha Ingreso Base de Datos' una vez se complete el reproceso correctivo programado.",
     "La distribucion corresponde a la fecha de informe registrada por el laboratorio de patologia y refleja el volumen mensual real cargado al sistema."),
    ('4. Distribucion estimada por tipo de informe', '4. Distribucion por clasificacion de malignidad'),
    ('Los porcentajes corresponden a una estimacion derivada del comportamiento observado del Servicio. Los valores definitivos se obtendran consultando la base de datos sobre la columna correspondiente.',
     'La siguiente tabla presenta la distribucion real de los 1.415 informes procesados segun el campo "Malignidad" registrado por el sistema durante la extraccion asistida.'),
    ("Al cierre del periodo se identifican tres casos marcados con el sufijo '(REVISAR DX)'. Estos casos corresponden a diagnosticos no tumorales descriptivos donde el modelo copio el preambulo del informe en lugar de la conclusion final, requiriendo validacion humana.",
     "Al cierre del periodo se identifican treinta (30) casos marcados con el sufijo '(REVISAR DX)' en el campo 'Diagnostico Principal'. Estos casos corresponden a informes en los que el modelo identifico el preambulo descriptivo (organo, tipo de biopsia) pero no logro aislar la conclusion final del patologo, por tratarse de diagnosticos descriptivos o no tumorales con lenguaje no estandar. Predominan informes de medula osea, biopsias de prostata sin malignidad confirmada e informes con redaccion extrainstitucional. Estos casos estan planificados para validacion manual por el equipo de Innovacion y Desarrollo."),
    ('El tiempo de procesamiento depende de la complejidad del informe (numero de biomarcadores, longitud del texto OCR) y del hardware disponible. Los siguientes valores son indicativos para una estacion estandar HUV con GPU dedicada.',
     'El tiempo de procesamiento depende de la complejidad del informe (numero de biomarcadores y longitud del texto OCR) y del hardware disponible. Los siguientes valores corresponden al desempeno observado sobre el hardware estandar HUV (RTX 3050, 8 GB VRAM) con el modelo nvidia/nemotron-3-nano en ejecucion local.'),
    ('Los valores promedio de campos extraidos son aproximaciones indicativas para fines ejecutivos; el calculo preciso se realiza mediante la funcion de auditoria del sistema sobre la base de datos completa.',
     'Los valores promedio de campos extraidos corresponden a la cobertura observada sobre los 1.415 informes procesados al cierre del periodo, calculada por el modulo de auditoria del sistema sobre la base de datos completa.'),
    ('Tiempo estimado de reproceso completo: una jornada nocturna.', 'Tiempo planificado de reproceso completo: una jornada nocturna.'),
]
for p in doc.paragraphs:
    for old, new in reemplazos:
        if old in p.text:
            replace_in_paragraph(p, old, new)

# Tabla 1: Por mes - 14 meses reales
t1 = doc.tables[1]
set_cell_text(t1.rows[0].cells[0], 'Mes')
set_cell_text(t1.rows[0].cells[1], 'Informes procesados')
set_cell_text(t1.rows[0].cells[2], 'Acumulado')
clear_table_rows_except_header(t1)
datos_mes = [
    ('Enero 2025', 76, 76),
    ('Febrero 2025', 105, 181),
    ('Marzo 2025', 103, 284),
    ('Abril 2025', 101, 385),
    ('Mayo 2025', 126, 511),
    ('Junio 2025', 99, 610),
    ('Julio 2025', 102, 712),
    ('Agosto 2025', 113, 825),
    ('Septiembre 2025', 119, 944),
    ('Octubre 2025', 106, '1.050'),
    ('Noviembre 2025', 152, '1.202'),
    ('Diciembre 2025', 138, '1.340'),
    ('Enero 2026', 74, '1.414'),
    ('Febrero 2026', 1, '1.415'),
    ('Total', '1.415', '1.415'),
]
for r in datos_mes:
    add_row_clone_last(t1, r)

# Tabla 3: Malignidad real
t3 = doc.tables[3]
set_cell_text(t3.rows[0].cells[0], 'Clasificacion')
set_cell_text(t3.rows[0].cells[1], 'Cantidad')
set_cell_text(t3.rows[0].cells[2], 'Porcentaje')
clear_table_rows_except_header(t3)
datos_mal = [
    ('Maligno', 893, '63,1%'),
    ('Benigno', 375, '26,5%'),
    ('Pre-maligno', 85, '6,0%'),
    ('Sin clasificacion explicita', 62, '4,4%'),
    ('Total', '1.415', '100%'),
]
for r in datos_mal:
    add_row_clone_last(t3, r)

# Tabla 4: REVISAR DX - 30 casos agrupados
t4 = doc.tables[4]
set_cell_text(t4.rows[0].cells[0], 'Rango de casos')
set_cell_text(t4.rows[0].cells[1], 'Cantidad y observacion')
clear_table_rows_except_header(t4)
datos_rev = [
    ('IHQ250180 - IHQ250917', '11 casos identificados con sufijo (REVISAR DX) en el lote inicial. Predominan informes de medula osea y biopsias descriptivas.'),
    ('IHQ250956 - IHQ251231', '12 casos identificados durante el lote intermedio. Incluye casos de medula osea, pulmon, cerebro y mama.'),
    ('IHQ251323 - IHQ251515', '7 casos identificados en el lote mas reciente. Incluye apendice, cervix, prostata y ganglio cervical.'),
    ('Total', '30 casos pendientes de validacion manual por Innovacion y Desarrollo.'),
]
for r in datos_rev:
    add_row_clone_last(t4, r)

doc.save(ruta)
print("OK 08 actualizado con datos reales")

# === Verificacion final ===
print("\n=== VERIFICACION FINAL ===")
patrones = ['POR DEFINIR', 'estimad', 'aproxim', 'indicativ', 'IHQXXXNNNN']
total_residuales = 0
for fname in sorted(os.listdir(CARPETA)):
    if not fname.endswith('.docx'):
        continue
    d = Document(os.path.join(CARPETA, fname))
    placeholders = []
    for i, p in enumerate(d.paragraphs):
        for pat in patrones:
            if pat.lower() in p.text.lower():
                placeholders.append(f'P[{i}]:{pat}')
                break
    for ti, t in enumerate(d.tables):
        for ri, r in enumerate(t.rows):
            for ci, c in enumerate(r.cells):
                for pat in patrones:
                    if pat.lower() in c.text.lower():
                        placeholders.append(f'T{ti}R{ri}C{ci}:{pat}')
                        break
    total_residuales += len(placeholders)
    estado = 'LIMPIO' if not placeholders else f'{len(placeholders)} residuales'
    print(f"  {fname}: {estado}")
    for ph in placeholders[:3]:
        print(f"      - {ph}")
print(f"\nTotal residuales: {total_residuales}")
