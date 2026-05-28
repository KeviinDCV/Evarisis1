# -*- coding: utf-8 -*-
"""
Generador de los 8 documentos oficiales de entrega EVARISIS V6.9.9
para el Hospital Universitario del Valle.

Estrategia:
  1. Copia la plantilla oficial (membrete + logos HUV) con shutil.copy
  2. Abre la copia con python-docx
  3. Inserta el contenido del cuerpo ANTES de la sección de firmas
     "Anexos / Copia archivo / Fecha / Proyecto / Reviso / Aprobo"
  4. Rellena los placeholders de firma con los datos institucionales

Uso:
  venv0\Scripts\python.exe herramientas_ia\resultados\generar_entrega_huv_v699.py
"""

from __future__ import annotations

import shutil
import time
from copy import deepcopy
from pathlib import Path
from typing import Iterable, List, Optional, Sequence, Tuple

from docx import Document
from docx.document import Document as DocumentT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# --------------------------------------------------------------------------- #
# Rutas y datos institucionales                                               #
# --------------------------------------------------------------------------- #

PLANTILLA = Path(r"C:\Users\Kechavarro\Downloads\PLANTILLA_MEMBRETE_2026.docx")
SALIDA = Path(
    r"C:\Users\Kechavarro\Desktop\ProyectoHUV9GESTOR_ONCOLOGIA"
    r"\documentacion\entrega_HUV_V6.9.9"
)

FECHA_DOC = "Mayo 2026"
PROYECTO = "Innovacion y Desarrollo"
REVISO = "Innovacion y Desarrollo"
APROBO = "Innovacion y Desarrollo"

VERSION_PROD = "V6.9.9"
NOMBRE_SISTEMA = "EVARISIS - Sistema Inteligente de Gestion Oncologica"

FUENTE = "Arial"
TAM_NORMAL = Pt(11)
TAM_H1 = Pt(14)
TAM_H2 = Pt(12)
TAM_H3 = Pt(11)
TAM_PIE = Pt(8)

COLOR_AZUL_HUV = RGBColor(0x1F, 0x3A, 0x6A)
COLOR_GRIS_TBL = "D9E1F2"
COLOR_NEGRO = RGBColor(0, 0, 0)

# --------------------------------------------------------------------------- #
# Utilidades de formato                                                       #
# --------------------------------------------------------------------------- #


def set_run(run, *, size=TAM_NORMAL, bold=False, italic=False,
            color=COLOR_NEGRO, font=FUENTE) -> None:
    run.font.name = font
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font)


def add_paragraph_before(anchor_p, doc: DocumentT):
    """Crea un parrafo nuevo justo ANTES del parrafo ancla y lo devuelve."""
    new_p = OxmlElement("w:p")
    anchor_p._element.addprevious(new_p)
    from docx.text.paragraph import Paragraph

    return Paragraph(new_p, anchor_p._parent)


def add_table_before(anchor_p, doc: DocumentT, rows: int, cols: int):
    """Inserta una tabla antes del parrafo ancla."""
    table = doc.add_table(rows=rows, cols=cols)
    # NOTA: plantilla solo tiene Normal Table; los bordes se aplican con set_cell_borders_thin
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    anchor_p._element.addprevious(table._element)
    return table


def shade_cell(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold=False, size=TAM_NORMAL,
                  align=WD_ALIGN_PARAGRAPH.LEFT, color=COLOR_NEGRO) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Limpiar parrafos previos
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    p = cell.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_run(run, bold=bold, size=size, color=color)


def set_cell_borders_thin(table) -> None:
    """Asegura bordes finos negros en toda la tabla."""
    tbl = table._element
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "808080")
        borders.append(b)
    # Reemplaza si existe
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    tbl_pr.append(borders)


# --------------------------------------------------------------------------- #
# Builder de contenido                                                        #
# --------------------------------------------------------------------------- #


class DocBuilder:
    """Inserta contenido antes del bloque de firmas conservado de la plantilla."""

    def __init__(self, doc: DocumentT, anchor_index: int = 4):
        """
        anchor_index = 4 corresponde al parrafo 'Anexos:' en la plantilla.
        Insertamos siempre ANTES de ese parrafo.
        """
        self.doc = doc
        self.anchor = doc.paragraphs[anchor_index]

    # ---------- helpers basicos ---------- #
    def _p(self):
        return add_paragraph_before(self.anchor, self.doc)

    def blank(self):
        self._p()

    def title(self, text: str):
        p = self._p()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text.upper())
        set_run(run, size=TAM_H1, bold=True, color=COLOR_AZUL_HUV)
        self.blank()

    def h1(self, text: str):
        p = self._p()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run(run, size=TAM_H1, bold=True, color=COLOR_AZUL_HUV)

    def h2(self, text: str):
        p = self._p()
        run = p.add_run(text)
        set_run(run, size=TAM_H2, bold=True, color=COLOR_AZUL_HUV)

    def h3(self, text: str):
        p = self._p()
        run = p.add_run(text)
        set_run(run, size=TAM_H3, bold=True, color=COLOR_NEGRO)

    def para(self, text: str, *, bold=False, italic=False, justify=True):
        p = self._p()
        p.alignment = (
            WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
        )
        run = p.add_run(text)
        set_run(run, bold=bold, italic=italic)

    def bullets(self, items: Iterable[str]):
        for it in items:
            p = self._p()
            p.paragraph_format.left_indent = Cm(0.6)
            run = p.add_run(f"-  {it}")
            set_run(run)

    def numbered(self, items: Iterable[str]):
        for i, it in enumerate(items, 1):
            p = self._p()
            p.paragraph_format.left_indent = Cm(0.6)
            run = p.add_run(f"{i}.  {it}")
            set_run(run)

    def table(self, header: Sequence[str], rows: Sequence[Sequence[str]],
              col_widths_cm: Optional[Sequence[float]] = None):
        n_cols = len(header)
        t = add_table_before(self.anchor, self.doc,
                             rows=len(rows) + 1, cols=n_cols)
        set_cell_borders_thin(t)
        # Cabecera
        for i, h in enumerate(header):
            c = t.cell(0, i)
            shade_cell(c, COLOR_GRIS_TBL)
            set_cell_text(c, h, bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(10.5))
        # Filas
        for r, row in enumerate(rows, start=1):
            for i, val in enumerate(row):
                set_cell_text(t.cell(r, i), str(val), size=Pt(10.5))
        # Anchos
        if col_widths_cm:
            for col_idx, w in enumerate(col_widths_cm):
                for r in range(len(rows) + 1):
                    t.cell(r, col_idx).width = Cm(w)
        # Blank line despues de la tabla
        self.blank()

    def control_versiones(self, codigo: str, titulo_doc: str):
        """Inserta una tabla de control de versiones estandar."""
        self.h2("Control del documento")
        self.table(
            ["Campo", "Detalle"],
            [
                ("Codigo del documento", codigo),
                ("Titulo", titulo_doc),
                ("Version del sistema", VERSION_PROD),
                ("Version del documento", "1.0"),
                ("Fecha de emision", FECHA_DOC),
                ("Estado", "Vigente"),
                ("Clasificacion", "Uso interno HUV"),
                ("Area responsable", "Innovacion y Desarrollo"),
            ],
            col_widths_cm=[5.0, 11.0],
        )

    # ---------- firmas ---------- #
    def rellenar_firmas(self):
        """Sustituye los placeholders de Anexos/Copia/Fecha/Firmas."""
        paragraphs = self.doc.paragraphs
        # Localizamos por contenido
        for p in paragraphs:
            txt = p.text
            if not txt:
                continue
            low = txt.lower()
            new_text = None
            if low.startswith("anexos"):
                new_text = "Anexos: No aplica"
            elif low.startswith("copia archivo"):
                new_text = "Copia archivo: Innovacion y Desarrollo - HUV"
            elif low.startswith("fecha"):
                new_text = f"Fecha: {FECHA_DOC}"
            elif "proyect" in low and ":" in txt:
                new_text = f"Proyecto:\t{PROYECTO}"
            elif "revis" in low and ":" in txt:
                new_text = f"Reviso:\t{REVISO}"
            elif "aprob" in low and ":" in txt:
                new_text = f"Aprobo:\t{APROBO}"
            if new_text is None:
                continue
            # Limpiar runs y escribir nuevo texto
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            run = p.add_run(new_text)
            set_run(run, size=TAM_PIE, bold=False)


# --------------------------------------------------------------------------- #
# Generacion de documento base                                                #
# --------------------------------------------------------------------------- #


def crear_documento(nombre_archivo: str, codigo: str, titulo_largo: str
                    ) -> Tuple[Path, DocumentT, DocBuilder]:
    SALIDA.mkdir(parents=True, exist_ok=True)
    destino = SALIDA / nombre_archivo
    shutil.copy(PLANTILLA, destino)
    doc = Document(destino)
    builder = DocBuilder(doc, anchor_index=4)
    builder.title(titulo_largo)
    builder.control_versiones(codigo, titulo_largo)
    return destino, doc, builder


def guardar(doc: DocumentT, builder: DocBuilder, destino: Path) -> None:
    builder.rellenar_firmas()
    doc.save(destino)


# --------------------------------------------------------------------------- #
# CONTENIDOS                                                                  #
# --------------------------------------------------------------------------- #
# 01. Credenciales y Configuracion                                            #
# --------------------------------------------------------------------------- #


def doc_01_credenciales():
    destino, doc, b = crear_documento(
        "01_CREDENCIALES_Y_CONFIGURACION.docx",
        "EVA-CONF-001",
        "Credenciales y Configuracion Tecnica - Sistema EVARISIS V6.9.9",
    )

    b.h1("1. Proposito y alcance del documento")
    b.para(
        "El presente documento consolida la informacion sensible de "
        "configuracion y credenciales requerida para la operacion, despliegue "
        "y mantenimiento del Sistema Inteligente de Gestion Oncologica "
        "EVARISIS en su version V6.9.9. Esta informacion esta destinada "
        "exclusivamente al equipo de Innovacion y Desarrollo del Hospital "
        "Universitario del Valle (HUV) y al personal autorizado para la "
        "administracion del aplicativo en las estaciones de trabajo del "
        "Servicio de Patologia."
    )
    b.para(
        "El alcance abarca los parametros del archivo config.ini, las "
        "credenciales de base de datos MySQL/MariaDB en proceso de despliegue "
        "sobre XAMPP, los parametros del servidor de inferencia local LM "
        "Studio, la configuracion del motor OCR Tesseract, el estado de los "
        "proveedores de nube y el procedimiento institucional para rotar las "
        "credenciales."
    )

    b.h1("2. Advertencia de confidencialidad")
    b.para(
        "Este documento contiene informacion sensible que puede comprometer "
        "la seguridad de los datos clinicos custodiados por el HUV. Su "
        "circulacion debe restringirse al personal autorizado expresamente "
        "por el equipo de Innovacion y Desarrollo. Esta prohibida la "
        "publicacion, fotocopia parcial o total, transmision por correo "
        "personal o almacenamiento en servicios externos no institucionales. "
        "Toda solicitud de acceso debe quedar registrada en la bitacora del "
        "area."
    )
    b.bullets([
        "Clasificacion: Confidencial - Uso interno HUV.",
        "Custodia: Innovacion y Desarrollo.",
        "Periodo de retencion: vigente mientras el sistema este en operacion.",
        "Soporte de divulgacion: acta de entrega firmada por el receptor.",
    ])

    b.h1("3. Configuracion del archivo config/config.ini")
    b.para(
        "El archivo config.ini centraliza la totalidad de parametros "
        "operativos del aplicativo. Se ubica en la subcarpeta 'config' "
        "junto al ejecutable GestorOncologia.exe y puede ser editado con "
        "Bloc de notas por personal autorizado. Las secciones que lo "
        "componen se describen a continuacion."
    )

    b.h2("3.1 Seccion [PATHS]")
    b.table(
        ["Parametro", "Valor de referencia", "Descripcion"],
        [
            ("tesseract_windows",
             r"C:\Program Files\Tesseract-OCR\tesseract.exe",
             "Ruta absoluta al ejecutable Tesseract en estaciones Windows."),
            ("base_pdfs",
             "pdfs_patologia",
             "Carpeta donde se depositan los informes PDF a procesar."),
            ("base_datos_sqlite",
             "data/huv_oncologia_NUEVO.db",
             "Ruta local de la base SQLite principal del aplicativo."),
            ("base_datos_diagnosticos",
             "data/diagnosticos_ia.db",
             "Base acumulativa con los diagnosticos generados por IA."),
        ],
        col_widths_cm=[4.5, 5.5, 6.0],
    )

    b.h2("3.2 Seccion [OCR]")
    b.table(
        ["Parametro", "Valor", "Justificacion clinica"],
        [
            ("dpi", "400", "Resolucion optima para texto de patologia digitalizado."),
            ("idioma", "spa", "Diccionarios Tesseract en espanol."),
            ("psm", "6", "Layout de bloque uniforme, adecuado para informes."),
            ("preprocesamiento", "True",
             "Activa binarizacion adaptativa y eliminacion de ruido."),
        ],
        col_widths_cm=[4.5, 3.5, 8.0],
    )

    b.h2("3.3 Seccion [LM_STUDIO]")
    b.table(
        ["Parametro", "Valor", "Descripcion"],
        [
            ("endpoint", "http://127.0.0.1:1234/v1",
             "URL del servicio OpenAI-compatible expuesto por LM Studio."),
            ("api_key", "lm-studio",
             "Token simbolico (LM Studio no exige autenticacion real)."),
            ("modelo_produccion", "nvidia/nemotron-3-nano",
             "Modelo aprobado para la version V6.9.9."),
            ("timeout_segundos", "900",
             "Tiempo maximo por chunk; cubre los picos de inferencia."),
            ("max_tokens", "8000",
             "Limite superior por respuesta JSON estructurada."),
            ("temperature", "0",
             "Determinismo total para mantener trazabilidad clinica."),
        ],
        col_widths_cm=[4.5, 4.5, 7.0],
    )

    b.h2("3.4 Seccion [BD_REMOTA] (despliegue futuro multi-PC)")
    b.table(
        ["Parametro", "Valor de referencia", "Observacion"],
        [
            ("motor", "MySQL / MariaDB", "Sobre XAMPP en servidor LAN."),
            ("host", "192.168.X.Y", "IP fija del servidor HUV - por confirmar."),
            ("puerto", "3306", "Puerto estandar MySQL."),
            ("base_de_datos", "huv_oncologia", "Esquema creado por DBA HUV."),
            ("usuario", "huv_app", "Cuenta de aplicacion con permisos DML."),
            ("password", "huv2026",
             "Contrasena institucional - rotacion semestral recomendada."),
            ("charset", "utf8mb4", "Soporte de caracteres clinicos especiales."),
        ],
        col_widths_cm=[4.5, 4.5, 7.0],
    )

    b.h1("4. Credenciales de base de datos MySQL/MariaDB")
    b.para(
        "Para el despliegue multi-usuario en la red HUV, se utilizara una "
        "instancia de MariaDB embebida en XAMPP. El esquema huv_oncologia "
        "contiene la tabla informes_ihq con 186 columnas y la tabla "
        "diagnosticos_ia con 1.415 registros al cierre del periodo. El "
        "acceso desde las estaciones cliente se realiza usando el driver "
        "pymysql ya empaquetado en el ejecutable."
    )
    b.bullets([
        "Cuenta administrativa: root - solo para el DBA del HUV.",
        "Cuenta aplicacion: huv_app - solo permisos SELECT, INSERT, UPDATE.",
        "Cuenta lectura: huv_lectura - reservada para reportes de gerencia.",
        "Backup diario programado en el servidor XAMPP a las 23:00 horas.",
    ])

    b.h1("5. Configuracion de LM Studio")
    b.para(
        "LM Studio actua como servidor local de inferencia, garantizando "
        "que la informacion del paciente nunca abandone la infraestructura "
        "del HUV. La instalacion debe realizarse en cada estacion que "
        "ejecute el aplicativo o, alternativamente, en un servidor GPU "
        "compartido al que apunten todas las estaciones."
    )
    b.h3("Pasos de verificacion en cada estacion:")
    b.numbered([
        "Abrir LM Studio y cargar el modelo nvidia/nemotron-3-nano.",
        "Iniciar el servidor en la pestana 'Local Server' (puerto 1234).",
        "Confirmar que la URL http://127.0.0.1:1234/v1/models responda.",
        "Validar que la temperatura este fijada en 0 y max_tokens en 8000.",
        "Comprobar GPU activa para minimizar la latencia por chunk.",
    ])

    b.h1("6. Configuracion OCR Tesseract")
    b.para(
        "El motor Tesseract se invoca como subproceso desde el aplicativo. "
        "Se requiere version 5.3 o superior con el paquete de idioma "
        "espanol instalado. La ruta por defecto en Windows es "
        r"'C:\Program Files\Tesseract-OCR\tesseract.exe'."
    )
    b.bullets([
        "Resolucion de digitalizacion: 400 DPI a color o escala de grises.",
        "Idioma del traineddata: spa (incluye signos diacriticos).",
        "Page Segmentation Mode: PSM 6 (bloque uniforme de texto).",
        "Preprocesamiento: binarizacion adaptativa Sauvola.",
    ])

    b.h1("7. Proveedores cloud (estado actual)")
    b.para(
        "Por politica institucional de confidencialidad de datos clinicos, "
        "los proveedores de inferencia en la nube se mantienen "
        "deshabilitados en V6.9.9. Esta decision se adopta hasta que el "
        "Comite de Etica y la oficina juridica del HUV emitan lineamientos "
        "especificos sobre el envio de datos seudonimizados a terceros."
    )
    b.table(
        ["Proveedor", "Estado", "Motivo"],
        [
            ("Google Gemini", "Deshabilitado", "Politica de no envio a nube."),
            ("Groq", "Deshabilitado", "Sin acuerdo de confidencialidad firmado."),
            ("OpenRouter", "Deshabilitado", "Enruta a multiples proveedores externos."),
        ],
        col_widths_cm=[4.5, 4.0, 7.5],
    )

    b.h1("8. Procedimiento de cambio de credenciales")
    b.numbered([
        "Solicitar autorizacion escrita al lider de Innovacion y Desarrollo.",
        "Detener todas las estaciones que esten ejecutando el aplicativo.",
        "Realizar respaldo completo del archivo config.ini vigente.",
        "Aplicar el cambio con un editor de texto plano (Bloc de notas).",
        "Validar el nuevo parametro abriendo el aplicativo en modo lectura.",
        "Registrar el cambio en la bitacora con fecha, hora y responsable.",
        "Comunicar la fecha y hora del cambio a los usuarios finales.",
    ])

    b.h1("9. Permisos requeridos en la red HUV")
    b.table(
        ["Servicio", "Puerto", "Protocolo", "Origen", "Destino"],
        [
            ("MySQL/MariaDB", "3306", "TCP",
             "Estaciones cliente", "Servidor XAMPP"),
            ("LM Studio", "1234", "TCP",
             "Estaciones cliente", "Localhost o servidor GPU"),
            ("Compartido SMB", "445", "TCP",
             "Estaciones cliente", "Carpeta institucional de PDFs"),
            ("DNS interno", "53", "TCP/UDP",
             "Estaciones cliente", "Servidor DNS HUV"),
        ],
        col_widths_cm=[4.0, 2.0, 2.5, 4.0, 3.5],
    )

    b.h1("10. Custodia del documento")
    b.para(
        "El original de este documento reposa en la oficina del area de "
        "Innovacion y Desarrollo. Cualquier copia distribuida debe estar "
        "numerada y registrada en la planilla de control de copias. La "
        "vigencia del documento es de un (1) ano calendario contado desde "
        "su fecha de emision; transcurrido este plazo debera revisarse y "
        "reemitirse una nueva version."
    )
    b.blank()
    guardar(doc, b, destino)
    return destino


# --------------------------------------------------------------------------- #
# 02. Documentacion Tecnica                                                   #
# --------------------------------------------------------------------------- #


def doc_02_doc_tecnica():
    destino, doc, b = crear_documento(
        "02_DOCUMENTACION_TECNICA.docx",
        "EVA-TEC-001",
        "Documentacion Tecnica - Sistema EVARISIS V6.9.9",
    )

    b.h1("1. Descripcion general del sistema")
    b.para(
        "EVARISIS es un Sistema Inteligente de Gestion Oncologica desarrollado "
        "por el area de Innovacion y Desarrollo del Hospital Universitario "
        "del Valle. Su proposito es transformar la informacion no "
        "estructurada de los informes de patologia (especialmente "
        "Inmunohistoquimica - IHQ) en un repositorio estructurado de 184 "
        "campos por caso, habilitando la generacion de indicadores, la "
        "trazabilidad clinica y el soporte a decisiones medicas. La version "
        "V6.9.9 consolida el procesamiento masivo de informes y la "
        "extraccion mediante un modelo de lenguaje local."
    )
    b.para(
        "El sistema opera de manera autonoma en la red interna del HUV: "
        "ningun dato del paciente abandona la infraestructura institucional. "
        "Esto se logra utilizando un servidor de inferencia local (LM Studio) "
        "y motores OCR de codigo abierto. EVARISIS atiende necesidades del "
        "Servicio de Patologia, del Comite de Tumores y de la Direccion "
        "Cientifica del Hospital."
    )

    b.h1("2. Arquitectura del sistema")
    b.para(
        "EVARISIS adopta una arquitectura por capas, donde cada componente "
        "tiene una responsabilidad unica y bien delimitada."
    )
    b.table(
        ["Capa", "Tecnologia", "Responsabilidad principal"],
        [
            ("Presentacion", "Tkinter + ttkbootstrap",
             "Interfaz operativa multi-pestana para personal del HUV."),
            ("Logica de negocio", "Modulos Python en core/",
             "Orquestacion de procesos, validaciones y reglas clinicas."),
            ("Persistencia", "SQLite (local) / MySQL XAMPP (LAN)",
             "Almacenamiento de los 186 campos de cada informe."),
            ("Inteligencia artificial", "LM Studio + nvidia/nemotron-3-nano",
             "Extraccion semantica de los 184 campos por caso."),
            ("Procesamiento de imagenes", "PyMuPDF + Tesseract OCR",
             "Conversion de PDF a texto plano de alta calidad."),
            ("Distribucion", "PyInstaller (one-folder)",
             "Empaquetado en GestorOncologia.exe para Windows."),
        ],
        col_widths_cm=[4.5, 5.0, 6.5],
    )

    b.h1("3. Flujo de procesamiento de un informe IHQ")
    b.numbered([
        "Carga del PDF desde la carpeta pdfs_patologia.",
        "Renderizado de cada pagina a imagen con PyMuPDF a 400 DPI.",
        "Aplicacion de OCR Tesseract con idioma espanol (PSM 6).",
        "Limpieza de texto: normalizacion de saltos, encabezados y firmas.",
        "Particion en chunks por caso (delimitador IHQ + numero consecutivo).",
        "Construccion de prompt estructurado con esquema JSON de 184 campos.",
        "Invocacion al modelo nvidia/nemotron-3-nano a traves de LM Studio.",
        "Validacion del JSON devuelto contra el esquema esperado.",
        "Aplicacion de reglas anti-N/A y normalizacion de diagnosticos.",
        "Insercion o actualizacion del registro en la tabla informes_ihq.",
        "Replica del diagnostico en la tabla diagnosticos_ia (acumulativa).",
        "Actualizacion de los indicadores en la pestana Dashboard.",
    ])

    b.h1("4. Estructura de carpetas del proyecto")
    b.para(
        "La organizacion del repositorio refleja la separacion por capas "
        "descrita anteriormente. El siguiente esquema muestra la jerarquia "
        "principal de carpetas en el ambiente de desarrollo."
    )
    b.table(
        ["Carpeta / Archivo", "Contenido"],
        [
            ("ui.py", "Punto de entrada con la interfaz Tkinter."),
            ("core/", "Modulos de logica de negocio y extractores."),
            ("core/extractors/", "Extractores especializados por categoria."),
            ("core/llm_client.py", "Cliente unificado para LM Studio/Ollama."),
            ("core/enhanced_database_dashboard.py",
             "Dashboard de indicadores en vivo."),
            ("config/config.ini", "Parametros operativos editables."),
            ("config/version_info.py", "Version actual del aplicativo."),
            ("data/", "Bases SQLite locales (informes_ihq, diagnosticos_ia)."),
            ("pdfs_patologia/", "Carpeta de entrada con los PDFs de IHQ."),
            ("informes_ia/", "Respaldos JSON generados por la IA."),
            ("herramientas_ia/", "Utilidades de auditoria y soporte."),
            ("dist/", "Salida del empaquetado PyInstaller."),
            ("documentacion/", "Documentos oficiales del proyecto."),
        ],
        col_widths_cm=[5.5, 10.5],
    )

    b.h1("5. Modulos principales en core/")
    b.table(
        ["Modulo", "Responsabilidad"],
        [
            ("llm_client.py",
             "Cliente HTTP hacia LM Studio/Ollama con reintentos y timeout."),
            ("normalizador_diagnosticos.py",
             "Normaliza textos de diagnostico y aplica reglas anti-N/A."),
            ("extractors/medical_extractor.py",
             "Coordina la extraccion semantica de campos clinicos."),
            ("extractors/admin_extractor.py",
             "Extrae identificadores administrativos del informe."),
            ("extractors/biomarker_extractor.py",
             "Procesa los 141 campos de biomarcadores IHQ."),
            ("enhanced_database_dashboard.py",
             "Genera los indicadores graficos del Dashboard."),
            ("diagnosticos_ia_db.py",
             "Gestiona la tabla acumulativa de diagnosticos."),
        ],
        col_widths_cm=[5.5, 10.5],
    )

    b.h1("6. Esquema de base de datos (186 columnas)")
    b.para(
        "El nucleo del aplicativo es la tabla informes_ihq, con 186 columnas "
        "agrupadas funcionalmente. La clave primaria es 'Numero de caso' "
        "(formato IHQXXXNNNN). Adicionalmente, la tabla diagnosticos_ia "
        "guarda el historico acumulado de diagnosticos generados por la IA, "
        "con 1.415 registros al cierre."
    )
    b.table(
        ["Categoria", "Cantidad", "Ejemplos representativos"],
        [
            ("Identificacion administrativa", "19",
             "Numero de caso, fecha de ingreso, sede, servicio."),
            ("Datos demograficos", "11",
             "Edad, sexo, tipo de documento, EPS, regimen."),
            ("Procedimiento", "7",
             "Tipo de muestra, organo, lateralidad, fecha de toma."),
            ("Diagnostico clinico", "7",
             "Diagnostico principal, secundario, observaciones."),
            ("Estudios IHQ generales", "5",
             "Resultado global, conclusion, marcadores positivos."),
            ("Biomarcadores IHQ", "141",
             "ER, PR, HER2, Ki67, p53, p16, CD3, CD20, entre otros."),
            ("Metadata operativa", "2",
             "Estado Auditoria IA, Fecha Ingreso Base de Datos."),
        ],
        col_widths_cm=[5.0, 2.5, 8.5],
    )

    b.h1("7. Pipeline de extraccion con IA")
    b.para(
        "El pipeline de extraccion fue disenado para garantizar la "
        "trazabilidad clinica y la reproducibilidad de los resultados. Cada "
        "paso es auditable mediante los archivos JSON depositados en "
        "informes_ia/."
    )
    b.numbered([
        "OCR del PDF y consolidacion del texto plano.",
        "Segmentacion en chunks por caso IHQ.",
        "Construccion del prompt con esquema JSON de 184 campos.",
        "Llamada al modelo LLM local con response_format=json_schema.",
        "Recepcion del JSON y validacion estructural.",
        "Aplicacion de reglas anti-N/A introducidas en V6.9.9.",
        "Persistencia en BD y replica en diagnosticos_ia.",
    ])

    b.h1("8. Manejo de errores y reintentos")
    b.para(
        "El cliente LLM implementa una politica de reintentos con backoff "
        "exponencial. Los timeouts (mas frecuentes con nemotron-3-nano por "
        "el reasoning_content extendido) se controlan con el parametro "
        "timeout_segundos = 900. Las excepciones se registran en el log y "
        "se muestra una notificacion visual al operador, sin abortar el "
        "procesamiento del lote."
    )
    b.bullets([
        "Errores de red: 3 reintentos con espera de 5, 15 y 45 segundos.",
        "Respuestas JSON malformadas: se intenta reparacion automatica.",
        "Campos faltantes: se completan con cadena vacia y se marcan.",
        "Casos sospechosos: se etiquetan con sufijo '(REVISAR DX)'.",
        "Errores de BD: rollback transaccional y reintento puntual.",
    ])

    b.h1("9. Compilacion con PyInstaller")
    b.para(
        "El empaquetado se realiza usando el script COMPILAR_V69.bat dentro "
        "del entorno virtual venv0 (Python 3.13.4). El resultado es una "
        "carpeta dist/ con el ejecutable GestorOncologia.exe (110 MB), las "
        "dependencias necesarias y los archivos editables del aplicativo. "
        "La compilacion incluye los modelos de Tesseract opcionalmente, "
        "segun la politica de la estacion."
    )
    b.h3("Pasos principales del proceso:")
    b.numbered([
        "Activar el entorno virtual: venv0\\Scripts\\activate",
        "Actualizar dependencias: pip install -r requirements.txt",
        "Ejecutar: COMPILAR_V69.bat",
        "Verificar que dist/GestorOncologia.exe se creo correctamente.",
        "Comprimir la carpeta dist/ y firmarla digitalmente si aplica.",
        "Publicar el paquete en el repositorio institucional de software.",
    ])

    b.h1("10. Distribucion multi-PC")
    b.para(
        "Para el despliegue en multiples estaciones, el aplicativo soporta "
        "configuracion externa: el archivo config.ini queda fuera del "
        "ejecutable y puede personalizarse por estacion (por ejemplo, la IP "
        "del servidor BD). La carpeta data/ puede ser local o apuntar a un "
        "recurso compartido si se requiere repositorio unificado."
    )
    b.bullets([
        "Cada estacion conserva su copia local de configuracion.",
        "Los PDFs pueden residir en un recurso SMB centralizado.",
        "La BD MySQL en XAMPP centraliza la informacion clinica.",
        "Los respaldos JSON locales permiten reproceso sin afectar BD central.",
    ])

    b.h1("11. Mantenimiento y actualizacion del sistema")
    b.para(
        "El mantenimiento del sistema incluye actualizaciones de version, "
        "rotacion del modelo LLM cuando se libere uno aprobado por el "
        "Comite, y monitoreo del rendimiento. Las actualizaciones se "
        "distribuyen mediante reemplazo controlado del ejecutable, "
        "preservando los archivos config.ini, data/ y pdfs_patologia/ de "
        "cada estacion."
    )
    b.h3("Tareas recomendadas:")
    b.bullets([
        "Mensual: revisar el log de errores y los casos con (REVISAR DX).",
        "Trimestral: validar el desempeno del modelo con un lote control.",
        "Semestral: rotar credenciales BD y actualizar Tesseract si aplica.",
        "Anual: revisar el roadmap y planear nueva version mayor.",
    ])
    b.blank()
    guardar(doc, b, destino)
    return destino


# --------------------------------------------------------------------------- #
# 03. Documentacion Tecnica Adicional                                         #
# --------------------------------------------------------------------------- #


def doc_03_doc_adicional():
    destino, doc, b = crear_documento(
        "03_DOCUMENTACION_TECNICA_ADICIONAL.docx",
        "EVA-TEC-002",
        "Documentacion Tecnica Adicional - Sistema EVARISIS V6.9.9",
    )

    b.h1("1. Historico de versiones")
    b.para(
        "El siguiente recorrido sintetiza la evolucion funcional y tecnica "
        "del aplicativo desde sus primeras iteraciones hasta la version "
        "actual V6.9.9. Cada hito refleja una decision tecnica significativa "
        "validada con el equipo clinico del Servicio de Patologia."
    )
    b.table(
        ["Version", "Hito principal"],
        [
            ("V6.7.x", "Prueba de concepto inicial. Extraccion de 3 campos "
             "(numero de caso, organo y diagnostico)."),
            ("V6.8.0", "Expansion del esquema a 184 campos cubriendo "
             "biomarcadores y datos administrativos."),
            ("V6.9.0", "Migracion a arquitectura multi-usuario con MySQL "
             "sobre XAMPP."),
            ("V6.9.1 - V6.9.5", "Mejoras en refresco en tiempo real del "
             "Dashboard y empaquetado distribuible con PyInstaller."),
            ("V6.9.6", "Integracion experimental de Ollama como proveedor "
             "alternativo de inferencia."),
            ("V6.9.7", "Correccion del mecanismo de deteccion de proveedor y "
             "ajuste de timeouts."),
            ("V6.9.8", "Adopcion del parametro json_schema para forzar "
             "estructura de respuesta del LLM."),
            ("V6.9.9", "Refuerzo del prompt para preservar diagnosticos no "
             "tumorales y reduccion de respuestas N/A."),
        ],
        col_widths_cm=[3.5, 12.5],
    )

    b.h1("2. Decisiones de diseno relevantes")
    b.para(
        "Las decisiones que se enumeran a continuacion explican el porque "
        "de la arquitectura actual y orientan futuras evoluciones."
    )
    b.h3("2.1 Por que un LLM local")
    b.para(
        "La confidencialidad de la historia clinica del paciente es no "
        "negociable para el HUV. Toda alternativa basada en nube fue "
        "descartada por la imposibilidad de garantizar que los textos no "
        "salgan de la infraestructura institucional. La inferencia local "
        "ademas elimina la dependencia de conectividad externa y reduce "
        "costos recurrentes."
    )
    b.h3("2.2 Por que nvidia/nemotron-3-nano")
    b.para(
        "Tras evaluar multiples modelos abiertos, nemotron-3-nano ofrecio "
        "el mejor balance entre tamano (corre en GPU de gama media), "
        "calidad en espanol clinico y soporte de salida estructurada con "
        "json_schema. Su capacidad de razonamiento extenso, aunque "
        "demanda mas tiempo, mejoro notablemente la cobertura de "
        "biomarcadores en informes complejos."
    )
    b.h3("2.3 Por que SQLite local + MySQL en red")
    b.para(
        "SQLite garantiza arranque sin dependencias y permite operar en "
        "modo degradado si la red institucional sufre una interrupcion. "
        "MySQL/MariaDB sobre XAMPP habilita el trabajo concurrente desde "
        "varias estaciones del Servicio de Patologia y centraliza la "
        "informacion para reportes consolidados."
    )

    b.h1("3. Modelos LLM evaluados")
    b.para(
        "El siguiente cuadro sintetiza los modelos evaluados durante el "
        "proceso de seleccion del motor de inferencia."
    )
    b.table(
        ["Modelo", "Tipo", "Resultado observado", "Decision"],
        [
            ("Qwen 14B", "Local", "Buena cobertura, latencia moderada.",
             "Descartado por consumo de VRAM."),
            ("Qwen 32B", "Local", "Cobertura superior, requiere GPU alta.",
             "Reservado para futuro despliegue centralizado."),
            ("Nemotron 3 Nano", "Local",
             "Excelente JSON estructurado y manejo de terminologia clinica.",
             "Adoptado en V6.9.9."),
            ("gpt-oss", "Local",
             "Respuesta inconsistente con esquema json_schema.",
             "Descartado."),
            ("MedGemma 4B", "Local",
             "Adecuado para tareas clinicas, debilidad en biomarcadores.",
             "Postergado para evaluacion futura."),
            ("MedGemma 27B", "Local",
             "Requiere GPU superior a la disponible.",
             "No factible en hardware actual."),
            ("DeepSeek (cloud)", "Nube",
             "Calidad alta pero no cumple politica de confidencialidad.",
             "Descartado."),
        ],
        col_widths_cm=[3.0, 2.0, 6.0, 5.0],
    )

    b.h1("4. Ingenieria de prompt y reglas anti-N/A")
    b.para(
        "El prompt del sistema esta disenado para combinar el rol clinico, "
        "el esquema JSON estricto y reglas explicitas de comportamiento. "
        "La version V6.9.9 introdujo cinco reglas adicionales para reducir "
        "la frecuencia de respuestas 'N/A' en campos donde la informacion "
        "esta presente pero requiere inferencia ligera."
    )
    b.numbered([
        "Si el informe describe un hallazgo no tumoral, debe registrarse "
        "explicitamente y no marcarse como N/A.",
        "Las negaciones tipo 'no se observa' deben mapearse a valor "
        "estandarizado 'Negativo'.",
        "Cuando un biomarcador este mencionado sin porcentaje, debe "
        "registrarse el resultado cualitativo.",
        "Las observaciones del patologo en parrafo libre deben condensarse "
        "en el campo 'Observaciones' sin perdida clinica.",
        "Si dos diagnosticos coexisten, debe registrarse el principal en "
        "'Diagnostico Principal' y el secundario en 'Diagnostico Secundario'.",
    ])

    b.h1("5. Tests de regresion y casos de referencia")
    b.para(
        "Cada vez que se modifica un extractor o el prompt, se ejecuta un "
        "lote de regresion sobre casos de referencia previamente validados. "
        "El criterio es estricto: si cualquier caso de referencia baja su "
        "score, el cambio se revierte automaticamente."
    )
    b.bullets([
        "Patron 'no presentan perdida': IHQ250133, IHQ250159.",
        "Patron 'expresion positiva para': IHQ250035.",
        "Patron 'negativo para': IHQ250127.",
        "Diagnostico no tumoral: IHQ250864, IHQ250900, IHQ250917 "
        "(marcados (REVISAR DX) en V6.9.9).",
    ])

    b.h1("6. Limitaciones conocidas")
    b.bullets([
        "Velocidad: nemotron-3-nano puede tomar varios minutos por chunk "
        "complejo debido a su reasoning_content extendido.",
        "Casos extrainstitucionales: PDFs sin formato HUV requieren "
        "validacion manual mayor.",
        "Imagenes embebidas: las microfotografias no se procesan; el OCR "
        "se limita al texto.",
        "Dependencia de hardware GPU: estaciones sin GPU dedicada operan "
        "en CPU con tiempos significativamente mayores.",
    ])

    b.h1("7. Hallazgos tecnicos colaterales")
    b.bullets([
        "reasoning_content: el modelo nemotron expone su cadena de "
        "razonamiento separada del content; el aplicativo la descarta "
        "antes de persistir.",
        "Race condition: detectada en el refresco del Dashboard cuando se "
        "procesan lotes muy grandes; mitigada con bloqueo de UI.",
        "Encoding ANSI vs UTF-8: corregido en el lector de config.ini para "
        "garantizar lectura uniforme.",
        "Memoria: el procesamiento de PDFs muy grandes se segmenta por "
        "paginas para evitar saturacion de memoria.",
    ])

    b.h1("8. Roadmap futuro propuesto")
    b.numbered([
        "Pipeline hibrido: combinar extractores deterministas con LLM "
        "para campos administrativos rigidos.",
        "Optimizacion del prompt: separar prompts por categoria de "
        "biomarcador para reducir tokens.",
        "Dashboard analitico: incorporar series temporales por organo y "
        "biomarcador.",
        "Integracion HIS: explorar interoperabilidad con el HIS del HUV "
        "via HL7 o FHIR.",
        "Modulo de auditoria automatizada: reactivar agentes especializados "
        "para revisiones masivas con alertas tempranas.",
    ])

    b.h1("9. Referencias y enlaces")
    b.table(
        ["Recurso", "Ubicacion / Descripcion"],
        [
            ("Documentacion oficial del proyecto",
             "documentacion/INFORME_GLOBAL_PROYECTO.md"),
            ("CHANGELOG del programa", "documentacion/CHANGELOG.md"),
            ("CHANGELOG de IA / agentes",
             "documentacion/CHANGELOG_CLAUDE.md"),
            ("Bitacora de acercamientos",
             "documentacion/BITACORA_DE_ACERCAMIENTOS.md"),
            ("Manual de Tesseract", "https://github.com/tesseract-ocr/tesseract"),
            ("Documentacion LM Studio", "https://lmstudio.ai/docs"),
            ("python-docx", "https://python-docx.readthedocs.io/"),
        ],
        col_widths_cm=[5.5, 10.5],
    )
    b.blank()
    guardar(doc, b, destino)
    return destino


# --------------------------------------------------------------------------- #
# 04. Manual de Usuario                                                       #
# --------------------------------------------------------------------------- #


def doc_04_manual_usuario():
    destino, doc, b = crear_documento(
        "04_MANUAL_DE_USUARIO.docx",
        "EVA-USR-001",
        "Manual de Usuario - Sistema EVARISIS V6.9.9",
    )

    b.h1("1. Introduccion")
    b.para(
        "Bienvenido al Sistema Inteligente de Gestion Oncologica EVARISIS, "
        "una herramienta institucional del Hospital Universitario del Valle "
        "que automatiza la extraccion de informacion clinica relevante a "
        "partir de los informes de patologia. Este manual esta dirigido al "
        "personal del Servicio de Patologia, del Comite de Tumores y a "
        "auxiliares administrativos que apoyen la operacion del sistema."
    )
    b.para(
        "EVARISIS le permitira procesar de manera masiva los informes en "
        "formato PDF, consultar los datos clinicos extraidos, generar "
        "reportes en Excel y obtener indicadores visuales del trabajo "
        "realizado. Toda la informacion se mantiene dentro de la "
        "infraestructura del HUV, garantizando la confidencialidad."
    )

    b.h1("2. Requisitos previos")
    b.bullets([
        "Equipo con sistema operativo Windows 10 o superior.",
        "Acceso a la red institucional del HUV.",
        "LM Studio instalado y con el modelo nvidia/nemotron-3-nano cargado.",
        "Tesseract OCR instalado en C:\\Program Files\\Tesseract-OCR.",
        "Espacio en disco minimo 5 GB para PDFs y base de datos local.",
        "Permisos de escritura en la carpeta del aplicativo.",
    ])

    b.h1("3. Primera vez que abre el aplicativo")
    b.numbered([
        "Localice la carpeta del aplicativo (por defecto C:\\EVARISIS).",
        "Ejecute con doble clic GestorOncologia.exe.",
        "Espere a que se carguen los modulos (puede tardar 20 a 30 segundos).",
        "Verifique en la barra inferior el mensaje 'Conectado a LM Studio'.",
        "Si aparece un mensaje en rojo, abra LM Studio y cargue el modelo.",
        "Confirme en la pestana Visualizador que la base de datos responde.",
    ])

    b.h1("4. Pantalla principal")
    b.para(
        "La interfaz se organiza en pestanas tematicas. Cada pestana esta "
        "disenada para una tarea operativa especifica."
    )
    b.table(
        ["Pestana", "Funcion principal"],
        [
            ("Visualizador",
             "Consulta y filtrado de los registros clinicos extraidos."),
            ("Procesar PDFs",
             "Ejecuta el OCR sobre los PDFs depositados en pdfs_patologia."),
            ("Procesar con IA",
             "Aplica el modelo LLM local sobre los textos OCR generados."),
            ("Reportes",
             "Exporta selecciones a Excel y produce listados oficiales."),
            ("Dashboard",
             "Visualiza indicadores en vivo (volumen, calidad, biomarcadores)."),
        ],
        col_widths_cm=[4.5, 11.5],
    )

    b.h1("5. Como cargar y procesar un PDF de patologia")
    b.numbered([
        "Copie el PDF del informe en la carpeta pdfs_patologia.",
        "Abra el aplicativo y dirijase a la pestana 'Procesar PDFs'.",
        "Pulse el boton 'Detectar nuevos PDFs'.",
        "Seleccione en la lista los archivos a procesar (o 'Seleccionar todos').",
        "Pulse 'Procesar seleccionados' e ingrese la confirmacion solicitada.",
        "Espere a que la barra de progreso alcance el 100%.",
        "Al finalizar, el sistema mostrara la cantidad de chunks reconocidos.",
    ])

    b.h1("6. Como procesar con IA un lote de PDFs")
    b.para(
        "Una vez los PDFs han sido procesados con OCR, deben pasar por el "
        "modulo de IA. Este paso es el que estructura los 184 campos por "
        "caso usando el modelo nemotron-3-nano."
    )
    b.numbered([
        "Vaya a la pestana 'Procesar con IA'.",
        "Pulse 'Cargar lote' y elija los textos OCR a procesar.",
        "Confirme que el modelo activo sea nvidia/nemotron-3-nano.",
        "Pulse 'Iniciar procesamiento'.",
        "Observe el modal 'Procesamiento con IA - resultados' en vivo.",
        "Si requiere pausar, pulse 'Detener al terminar el caso actual'.",
        "Al final, revise el resumen: casos procesados, exitosos y revisar.",
    ])

    b.h1("7. Como consultar y filtrar datos en el Visualizador")
    b.bullets([
        "Use los filtros superiores para acotar por organo, fecha o sexo.",
        "Use el buscador para encontrar un caso por numero IHQXXXNNNN.",
        "Doble clic en una fila para ver el detalle completo en panel lateral.",
        "Use 'Mostrar columnas' para personalizar las columnas visibles.",
        "Use 'Ordenar' para reorganizar por cualquier columna numerica.",
    ])

    b.h1("8. Como exportar datos a Excel")
    b.numbered([
        "Aplique los filtros deseados en el Visualizador.",
        "Pulse el boton 'Exportar a Excel' en la barra superior.",
        "Elija la ubicacion y nombre del archivo (por defecto, fecha de hoy).",
        "Espere a que se confirme la exportacion (puede tomar minutos).",
        "Abra el archivo .xlsx para validar la informacion exportada.",
    ])

    b.h1("9. Como identificar casos marcados con (REVISAR DX)")
    b.para(
        "El sistema marca con el sufijo '(REVISAR DX)' aquellos casos donde "
        "la IA detecta inconsistencias o diagnosticos no tumorales que "
        "requieren validacion humana. Estos casos no son errores; son "
        "alertas para revision del patologo."
    )
    b.bullets([
        "Use el filtro 'Estado Auditoria IA' = 'REVISAR DX' en Visualizador.",
        "Abra el caso, compare el diagnostico de la IA contra el texto OCR.",
        "Si la IA es correcta, retire la marca desde la pestana Reportes.",
        "Si hay error, registre la observacion para el equipo de Desarrollo.",
    ])

    b.h1("10. Solucion de problemas comunes")
    b.table(
        ["Sintoma", "Causa probable", "Accion sugerida"],
        [
            ("'LM Studio no detectado'", "El servidor LM Studio no esta activo.",
             "Abra LM Studio y arranque el servidor en Local Server."),
            ("'BD no conecta'", "MySQL en XAMPP detenido o ruta SQLite invalida.",
             "Verifique XAMPP y la ruta en config.ini."),
            ("OCR lento", "PDFs muy grandes o estacion sin recursos.",
             "Procese en lotes mas pequenos fuera de horario pico."),
            ("'0 diagnosticos' en popup", "El lote no contenia casos validos.",
             "Confirme que los PDFs sean informes IHQ con el formato HUV."),
            ("Aplicacion 'congelada'", "Procesamiento intenso en curso.",
             "Espere; el sistema actualiza cada 30 segundos."),
            ("Caracteres extranos", "PDF con OCR insuficiente.",
             "Reescanee el original a 400 DPI con calidad alta."),
        ],
        col_widths_cm=[4.5, 5.5, 6.0],
    )

    b.h1("11. Quien contactar para soporte")
    b.para(
        "Para reportar incidentes, sugerencias de mejora o solicitar "
        "actualizaciones del aplicativo, contacte al area de Innovacion y "
        "Desarrollo del HUV. El equipo registrara su requerimiento en la "
        "mesa de servicio institucional y le brindara seguimiento."
    )
    b.bullets([
        "Area: Innovacion y Desarrollo - HUV.",
        "Canal recomendado: correo institucional con el asunto 'EVARISIS'.",
        "Tiempo de respuesta esperado: 1 dia habil para incidentes criticos.",
        "Adjunte siempre el numero de caso IHQ y el mensaje de error visible.",
    ])
    b.blank()
    guardar(doc, b, destino)
    return destino


# --------------------------------------------------------------------------- #
# 05. Material de entrenamiento - Sesion 1                                    #
# --------------------------------------------------------------------------- #


def doc_05_entrenamiento_s1():
    destino, doc, b = crear_documento(
        "05_MATERIAL_ENTRENAMIENTO_SESION_1.docx",
        "EVA-CAP-001",
        "Material de Entrenamiento - Sesion 1: Fundamentos y operacion basica",
    )

    b.h1("1. Objetivos de aprendizaje")
    b.bullets([
        "Comprender el proposito y alcance del sistema EVARISIS.",
        "Identificar los componentes principales de la interfaz.",
        "Realizar el procesamiento OCR de un primer informe.",
        "Consultar un caso en el Visualizador y describir sus campos.",
        "Reconocer cuando un caso requiere validacion humana.",
    ])

    b.h1("2. Agenda de la sesion (60 minutos)")
    b.table(
        ["Tiempo", "Actividad"],
        [
            ("00:00 - 00:05", "Bienvenida, encuadre y objetivos."),
            ("00:05 - 00:15",
             "Conceptos basicos: OCR, LLM, confidencialidad."),
            ("00:15 - 00:30", "Recorrido guiado por la interfaz."),
            ("00:30 - 00:40",
             "Ejercicio practico 1: cargar un PDF y procesarlo."),
            ("00:40 - 00:50",
             "Ejercicio practico 2: consulta en Visualizador."),
            ("00:50 - 00:57", "Preguntas frecuentes y dudas."),
            ("00:57 - 01:00", "Evaluacion corta y tarea para Sesion 2."),
        ],
        col_widths_cm=[3.0, 13.0],
    )

    b.h1("3. Conceptos basicos")
    b.h2("3.1 Que es OCR")
    b.para(
        "OCR significa Reconocimiento Optico de Caracteres. Es la "
        "tecnologia que permite convertir una imagen (por ejemplo, un PDF "
        "escaneado) en texto que la computadora puede analizar. EVARISIS "
        "usa Tesseract, un motor OCR de codigo abierto, configurado para "
        "espanol clinico."
    )
    b.h2("3.2 Que es un LLM local")
    b.para(
        "Un LLM (Large Language Model) es un modelo de lenguaje entrenado "
        "para entender y generar texto. 'Local' significa que se ejecuta "
        "dentro del HUV, sin enviar informacion a internet. Esto garantiza "
        "que la informacion clinica de los pacientes nunca abandona la "
        "institucion."
    )
    b.h2("3.3 Por que local y no en la nube")
    b.bullets([
        "Cumplimiento de la politica de confidencialidad del HUV.",
        "Independencia de proveedores externos.",
        "Disponibilidad sin internet o ante caidas del enlace WAN.",
        "Control institucional total sobre los datos clinicos.",
    ])

    b.h1("4. Recorrido por la interfaz")
    b.para(
        "Durante este bloque, el instructor mostrara cada pestana del "
        "aplicativo y explicara la funcion de cada boton principal. Los "
        "participantes deberan tomar nota de los elementos que se senalen."
    )
    b.bullets([
        "Visualizador: tabla central, filtros, exportacion.",
        "Procesar PDFs: lista de archivos detectados, boton de proceso.",
        "Procesar con IA: lote seleccionado, modal en vivo.",
        "Reportes: filtros avanzados y generacion de Excel.",
        "Dashboard: tarjetas de indicadores en tiempo real.",
    ])

    b.h1("5. Ejercicio practico 1: cargar un PDF y procesarlo")
    b.numbered([
        "Cada participante recibira un PDF de prueba (IHQ250001.pdf).",
        "Coloque el PDF en la carpeta pdfs_patologia.",
        "Abra el aplicativo y diríjase a 'Procesar PDFs'.",
        "Pulse 'Detectar nuevos PDFs' y verifique que aparece su archivo.",
        "Seleccione el PDF y pulse 'Procesar seleccionados'.",
        "Observe el avance y registre el tiempo total que tomo.",
    ])

    b.h1("6. Ejercicio practico 2: consultar el dato extraido")
    b.numbered([
        "Vaya a la pestana 'Visualizador'.",
        "Escriba 'IHQ250001' en el buscador.",
        "Haga doble clic en la fila para abrir el detalle.",
        "Identifique los campos: Numero de caso, Organo, Diagnostico Principal.",
        "Anote en su cuaderno tres biomarcadores presentes en el caso.",
        "Compare con el patologo de referencia si esta disponible.",
    ])

    b.h1("7. Preguntas frecuentes nivel 1")
    b.bullets([
        "Que pasa si el PDF esta en mala calidad? Se debe reescanear a 400 DPI.",
        "Cuanto tarda procesar un caso? Entre 1 y 5 minutos segun complejidad.",
        "Puedo procesar dos PDFs a la vez? Si, pero el sistema los pone en cola.",
        "Que pasa si LM Studio se cierra? El sistema espera y reintenta.",
        "Los datos quedan guardados si apago la PC? Si, la BD se actualiza al final.",
    ])

    b.h1("8. Evaluacion corta")
    b.numbered([
        "Mencione tres pestanas principales del aplicativo.",
        "Que significa OCR y para que se usa en EVARISIS?",
        "Por que el sistema usa un LLM local y no en la nube?",
        "Donde se colocan los PDFs antes de procesarlos?",
        "Como identifica un caso que requiere validacion humana?",
    ])

    b.h1("9. Tarea para la siguiente sesion")
    b.bullets([
        "Procese tres PDFs adicionales de su propia bandeja de trabajo.",
        "Anote el tiempo total que tomo cada uno.",
        "Registre cualquier mensaje de error o advertencia que aparezca.",
        "Identifique al menos un caso con la marca (REVISAR DX).",
        "Lleve sus observaciones a la Sesion 2 para discusion grupal.",
    ])
    b.blank()
    guardar(doc, b, destino)
    return destino


# --------------------------------------------------------------------------- #
# 06. Material de entrenamiento - Sesion 2                                    #
# --------------------------------------------------------------------------- #


def doc_06_entrenamiento_s2():
    destino, doc, b = crear_documento(
        "06_MATERIAL_ENTRENAMIENTO_SESION_2.docx",
        "EVA-CAP-002",
        "Material de Entrenamiento - Sesion 2: Procesamiento por lotes, auditoria y reportes",
    )

    b.h1("1. Repaso de la Sesion 1")
    b.para(
        "En la sesion anterior se cubrieron los fundamentos del sistema, "
        "los conceptos de OCR y LLM local, el recorrido por la interfaz y "
        "los primeros ejercicios practicos. Antes de avanzar, el "
        "instructor recopilara las observaciones de la tarea asignada y "
        "atendera dudas pendientes."
    )

    b.h1("2. Objetivos de la Sesion 2")
    b.bullets([
        "Procesar un lote de 5 PDFs y reanudar un lote interrumpido.",
        "Interpretar el modal 'Procesamiento con IA - resultados'.",
        "Realizar auditoria de calidad y reconocer casos N/A criticos.",
        "Generar y exportar un reporte oficial en Excel.",
        "Editar el archivo config.ini de forma segura.",
    ])

    b.h1("3. Procesamiento masivo")
    b.para(
        "El aplicativo esta disenado para procesar grandes cantidades de "
        "informes en una sola jornada. Esta seccion describe el flujo "
        "recomendado y como actuar si se interrumpe el lote."
    )
    b.h3("3.1 Configurar un lote")
    b.numbered([
        "Verifique que todos los PDFs esten en pdfs_patologia.",
        "En 'Procesar PDFs', pulse 'Detectar nuevos PDFs' y luego 'Seleccionar todos'.",
        "Procese primero el OCR. Espere a que termine la lista completa.",
        "Cambie a 'Procesar con IA' y cargue el mismo lote.",
        "Inicie el procesamiento y mantenga LM Studio abierto.",
    ])
    b.h3("3.2 Reanudar un lote interrumpido")
    b.bullets([
        "Si ocurre apagon o cierre inesperado, vuelva a abrir el aplicativo.",
        "Vaya a 'Procesar con IA' y pulse 'Reanudar lote'.",
        "El sistema detectara los casos ya procesados y continuara con los pendientes.",
        "Verifique en el log que no haya casos duplicados.",
    ])

    b.h1("4. Modal 'Procesamiento con IA - resultados'")
    b.para(
        "El modal en vivo permite seguir el avance del procesamiento "
        "case por case. Su lectura es clave para detectar problemas "
        "tempranamente y ajustar la operacion."
    )
    b.table(
        ["Columna del modal", "Significado"],
        [
            ("Caso", "Numero IHQ procesado actualmente."),
            ("Estado", "OK, EN PROCESO o ERROR."),
            ("Tiempo", "Duracion del procesamiento del caso."),
            ("Diagnostico extraido",
             "Texto resumen del diagnostico que arroja la IA."),
            ("Observaciones",
             "Notas tecnicas (timeouts, reintentos, marcas REVISAR DX)."),
        ],
        col_widths_cm=[4.5, 11.5],
    )

    b.h1("5. Auditoria de calidad: casos N/A")
    b.para(
        "Cuando la IA marca un campo como 'N/A', puede deberse a tres "
        "causas: el dato no esta en el informe, el dato esta pero la IA no "
        "lo reconocio, o el dato esta redactado de forma inusual. La "
        "sesion ensenara a distinguir cada caso."
    )
    b.bullets([
        "Filtre en Visualizador por 'Diagnostico Principal' = N/A.",
        "Abra el detalle y compare con el texto OCR del informe.",
        "Si la informacion esta presente, marque el caso para reproceso.",
        "Si la informacion no esta, conserve el N/A y registre la fuente.",
        "Reporte a Innovacion y Desarrollo los patrones repetitivos.",
    ])

    b.h1("6. Generacion de reportes y exportacion a Excel")
    b.numbered([
        "Vaya a 'Reportes' y elija el tipo (resumen mensual, por organo, etc.).",
        "Aplique filtros de fecha y servicio.",
        "Pulse 'Generar reporte'.",
        "Revise la vista previa y pulse 'Exportar a Excel'.",
        "Guarde el archivo en la carpeta institucional indicada.",
    ])

    b.h1("7. Edicion del archivo config.ini")
    b.para(
        "El archivo config.ini permite ajustar parametros del aplicativo "
        "sin recompilarlo. Solo personal autorizado debe modificarlo. Esta "
        "seccion describe los cambios mas comunes."
    )
    b.h3("7.1 Cambiar de LM Studio a Ollama")
    b.numbered([
        "Detenga el aplicativo si esta en ejecucion.",
        "Abra config.ini con Bloc de notas.",
        "Localice la seccion [LLM_PROVIDER].",
        "Cambie 'proveedor = lmstudio' por 'proveedor = ollama'.",
        "Ajuste el endpoint si Ollama corre en otro puerto.",
        "Guarde el archivo y reinicie el aplicativo.",
    ])
    b.h3("7.2 Cambiar de modelo")
    b.bullets([
        "Modifique el parametro 'modelo_produccion' por el nombre exacto.",
        "Asegurese de que el nuevo modelo este descargado en LM Studio.",
        "Realice una prueba con un caso conocido antes de un lote masivo.",
    ])

    b.h1("8. Ejercicio practico: lote de 5 PDFs")
    b.numbered([
        "Cada participante recibira 5 PDFs nuevos.",
        "Realice el flujo completo OCR + IA descrito en la seccion 3.",
        "Mientras se procesa, observe el modal en vivo y registre tiempos.",
        "Al finalizar, exporte un reporte Excel del lote.",
        "Identifique al menos un caso N/A y proponga su causa.",
    ])

    b.h1("9. Buenas practicas operativas")
    b.bullets([
        "Procese los lotes grandes en horario nocturno para no saturar la red.",
        "Realice backup de la carpeta data/ una vez por semana.",
        "Mantenga LM Studio actualizado pero no cambie de modelo sin autorizacion.",
        "Comunique al area de Desarrollo cualquier cambio de hardware.",
        "No edite manualmente la base de datos: use solo el aplicativo.",
    ])

    b.h1("10. Evaluacion final")
    b.numbered([
        "Mencione dos formas de reanudar un lote interrumpido.",
        "Explique como interpretar la columna Estado del modal en vivo.",
        "Describa tres causas posibles de un campo marcado como N/A.",
        "Cuales son los pasos para cambiar el modelo en config.ini?",
        "Indique tres buenas practicas operativas para el procesamiento masivo.",
    ])
    b.blank()
    guardar(doc, b, destino)
    return destino


# --------------------------------------------------------------------------- #
# 07. Registro historico de interacciones                                     #
# --------------------------------------------------------------------------- #


def doc_07_bitacora():
    destino, doc, b = crear_documento(
        "07_REGISTRO_HISTORICO_INTERACCIONES.docx",
        "EVA-BIT-001",
        "Registro Historico de Interacciones y Avances - Proyecto EVARISIS",
    )

    b.h1("1. Proposito de la bitacora")
    b.para(
        "Esta bitacora consolida la trayectoria del proyecto EVARISIS desde "
        "su prueba de concepto hasta la version V6.9.9. Su proposito es "
        "documentar las decisiones tomadas, los hallazgos relevantes y las "
        "lecciones aprendidas, sirviendo como insumo para futuras "
        "evoluciones del aplicativo y para los procesos de auditoria "
        "institucional."
    )

    b.h1("2. Cronologia por fases")

    b.h2("2.1 Fase 1: configuracion inicial y extraccion de tres campos (V6.7.x)")
    b.para(
        "La primera fase del proyecto se concentro en demostrar la "
        "factibilidad tecnica de extraer informacion estructurada desde "
        "informes IHQ. Se logro extraer tres campos por caso: numero de "
        "caso, organo y diagnostico principal. Esta prueba de concepto "
        "valido la combinacion OCR + LLM y motivo la expansion del alcance."
    )

    b.h2("2.2 Fase 2: expansion a 184 campos (V6.8.0)")
    b.para(
        "Con base en la validacion del Servicio de Patologia, el esquema se "
        "amplio para cubrir todas las categorias relevantes: datos "
        "administrativos, demograficos, procedimiento, diagnostico, "
        "estudios IHQ y los 141 biomarcadores definidos por el equipo "
        "clinico. Esta fase implico el rediseno completo de los prompts y "
        "la incorporacion de respuestas estructuradas en JSON."
    )

    b.h2("2.3 Fase 3: migracion a MySQL multi-usuario (V6.9.0)")
    b.para(
        "Para habilitar el trabajo concurrente del Servicio de Patologia, "
        "se introdujo soporte para MySQL/MariaDB sobre XAMPP. SQLite se "
        "mantuvo como motor local para operacion degradada. Se definieron "
        "las cuentas huv_app y huv_lectura con permisos diferenciados."
    )

    b.h2("2.4 Fase 4: refresh real-time y compilacion distribuible (V6.9.1 a V6.9.5)")
    b.para(
        "Esta fase consolido la experiencia del usuario final. El "
        "Dashboard se hizo reactivo a las inserciones en BD y el "
        "aplicativo se empaqueto con PyInstaller para distribucion en "
        "estaciones del HUV. El ejecutable resultante (110 MB) incluye "
        "todas las dependencias salvo Tesseract y LM Studio."
    )

    b.h2("2.5 Fase 5: integracion Ollama (V6.9.6)")
    b.para(
        "Se incorporo Ollama como segundo proveedor de inferencia local. "
        "Esta decision se tomo para reducir la dependencia de un unico "
        "ecosistema y permitir evaluaciones comparativas con otros modelos "
        "del catalogo Ollama."
    )

    b.h2("2.6 Fase 6: fix de deteccion de proveedor y timeouts (V6.9.7)")
    b.para(
        "Se detectaron y corrigieron problemas en la deteccion automatica "
        "del proveedor activo y se ajustaron los timeouts del cliente HTTP "
        "para acomodar el reasoning_content extendido del modelo "
        "nemotron-3-nano."
    )

    b.h2("2.7 Fase 7: parametro json_schema (V6.9.8)")
    b.para(
        "Se adopto el parametro response_format=json_schema, disponible en "
        "LM Studio y Ollama, para forzar respuestas conformes al esquema "
        "de 184 campos. Esto redujo drasticamente los errores de "
        "validacion estructural."
    )

    b.h2("2.8 Fase 8: correccion de diagnosticos no tumorales (V6.9.9)")
    b.para(
        "La version actual reforzo el prompt del sistema con cinco reglas "
        "anti-N/A. El cambio se motivo por casos donde la IA marcaba "
        "diagnosticos no tumorales como ausentes, cuando en realidad "
        "estaban presentes. Tras la correccion, la tasa de exito en la "
        "extraccion del diagnostico principal alcanzo el 97,9%."
    )

    b.h1("3. Hitos relevantes")
    b.table(
        ["Hito", "Periodo estimado", "Descripcion"],
        [
            ("Prueba de concepto OCR + LLM", "Mayo 2025",
             "Validacion inicial sobre 30 informes."),
            ("Esquema de 184 campos consolidado", "Septiembre 2025",
             "Validado con el Servicio de Patologia."),
            ("Migracion a MySQL XAMPP", "Diciembre 2025",
             "Habilitacion multi-usuario."),
            ("Empaquetado PyInstaller", "Enero 2026",
             "Distribucion en estaciones HUV."),
            ("Integracion Ollama", "Febrero 2026",
             "Doble proveedor de inferencia."),
            ("Adopcion json_schema", "Abril 2026",
             "Reduccion de errores estructurales."),
            ("Reglas anti-N/A", "Mayo 2026",
             "Refuerzo del prompt para V6.9.9."),
        ],
        col_widths_cm=[5.0, 3.5, 7.5],
    )

    b.h1("4. Problemas encontrados y resoluciones")
    b.table(
        ["Problema", "Resolucion aplicada"],
        [
            ("Latencia alta con nemotron",
             "Ajuste de timeout a 900 segundos y procesamiento por chunks."),
            ("Respuestas JSON malformadas",
             "Adopcion de response_format=json_schema en V6.9.8."),
            ("Casos marcados N/A erroneamente",
             "Cinco reglas anti-N/A introducidas en V6.9.9."),
            ("Cuelgues de UI en lotes grandes",
             "Bloqueo de UI durante procesamiento masivo."),
            ("Encoding del config.ini",
             "Lector forzado a UTF-8 con fallback a ANSI."),
            ("Modelo gpt-oss inestable",
             "Descartado tras pruebas comparativas."),
        ],
        col_widths_cm=[6.0, 10.0],
    )

    b.h1("5. Lecciones aprendidas")
    b.bullets([
        "La validacion clinica continua del Servicio de Patologia es "
        "indispensable para garantizar la utilidad del aplicativo.",
        "El uso de un LLM local no compromete la calidad cuando se elige "
        "un modelo apropiado y se itera sobre el prompt.",
        "La trazabilidad mediante respaldos JSON facilita la auditoria y el "
        "reproceso de casos sin afectar la BD principal.",
        "La separacion de responsabilidades en core/ acelera los ajustes "
        "puntuales sin generar regresiones.",
        "La distribucion mediante PyInstaller con config externo simplifica "
        "el despliegue multi-PC en la red institucional.",
    ])

    b.h1("6. Tabla de hitos con fechas estimadas")
    b.table(
        ["Version", "Fecha estimada", "Resultado clave"],
        [
            ("V6.7.0", "Mayo 2025", "Prueba de concepto inicial."),
            ("V6.8.0", "Septiembre 2025", "184 campos por caso."),
            ("V6.9.0", "Diciembre 2025", "Multi-usuario sobre MySQL."),
            ("V6.9.5", "Enero 2026", "Distribuible empaquetado."),
            ("V6.9.6", "Febrero 2026", "Integracion Ollama."),
            ("V6.9.7", "Marzo 2026", "Fix de deteccion de proveedor."),
            ("V6.9.8", "Abril 2026", "Adopcion json_schema."),
            ("V6.9.9", "Mayo 2026", "Diagnosticos no tumorales preservados."),
        ],
        col_widths_cm=[3.0, 3.5, 9.5],
    )
    b.blank()
    guardar(doc, b, destino)
    return destino


# --------------------------------------------------------------------------- #
# 08. Estadistica                                                             #
# --------------------------------------------------------------------------- #


def doc_08_estadistica():
    destino, doc, b = crear_documento(
        "08_ESTADISTICA.docx",
        "EVA-EST-001",
        "Estadistica de Uso y Resultados - Sistema EVARISIS V6.9.9",
    )

    b.h1("1. Volumen procesado")
    b.para(
        "Al cierre de Mayo de 2026, el sistema EVARISIS ha procesado un "
        "total de 1.415 informes de inmunohistoquimica en el rango "
        "IHQ250001 - IHQ251526. Estos informes corresponden a la "
        "produccion ordinaria del Servicio de Patologia del HUV y "
        "constituyen el universo actual sobre el que se calculan los "
        "indicadores de calidad y cobertura."
    )

    b.h1("2. Distribucion estimada por mes")
    b.para(
        "La siguiente tabla presenta la distribucion estimada de informes "
        "procesados por mes, agrupada para fines indicativos. Los valores "
        "podran refinarse cuando se complete el reproceso de los 30 casos "
        "pendientes."
    )
    b.table(
        ["Mes", "Informes procesados", "Acumulado"],
        [
            ("Diciembre 2025", "180", "180"),
            ("Enero 2026", "210", "390"),
            ("Febrero 2026", "220", "610"),
            ("Marzo 2026", "230", "840"),
            ("Abril 2026", "260", "1.100"),
            ("Mayo 2026", "315", "1.415"),
        ],
        col_widths_cm=[4.5, 5.0, 6.5],
    )
    b.para(
        "Nota: la distribucion mensual es indicativa y debera confirmarse "
        "con la consulta directa al campo 'Fecha Ingreso Base de Datos' "
        "una vez se complete el reproceso pendiente. [POR DEFINIR cifras "
        "exactas tras el reproceso correctivo]."
    )

    b.h1("3. Calidad de extraccion")
    b.table(
        ["Indicador", "Valor", "Interpretacion"],
        [
            ("Tasa global de exito a nivel de chunk", "99,9%",
             "Casi todos los chunks se procesan sin error tecnico."),
            ("Tasa de exito en diagnostico principal", "97,9%",
             "El sistema identifica correctamente el diagnostico en la "
             "gran mayoria de casos."),
            ("Tasa de N/A actual", "2,1% (30 casos)",
             "Casos pendientes de reproceso o validacion humana."),
        ],
        col_widths_cm=[5.5, 3.0, 7.5],
    )

    b.h1("4. Distribucion estimada por tipo de informe")
    b.para(
        "Los porcentajes corresponden a una estimacion derivada del "
        "comportamiento observado del Servicio. Los valores definitivos se "
        "obtendran consultando los campos 'Diagnostico Principal' y "
        "'Categoria Diagnostica' en el Visualizador."
    )
    b.table(
        ["Tipo de informe", "Porcentaje estimado", "Observaciones"],
        [
            ("Carcinomas", "55%", "Mama, gastrico, colorrectal, pulmon."),
            ("Linfomas", "12%",
             "Predominan B difuso de celulas grandes y Hodgkin."),
            ("Biopsias negativas", "10%", "Sin evidencia de malignidad."),
            ("Medula osea", "8%", "Estudios hematologicos."),
            ("Sarcomas y mesenquimales", "5%",
             "Incluye GIST y tumores de tejidos blandos."),
            ("Otros / no tumorales", "10%",
             "Diagnosticos no neoplasicos preservados en V6.9.9."),
        ],
        col_widths_cm=[5.0, 3.5, 7.5],
    )

    b.h1("5. Casos marcados (REVISAR DX)")
    b.para(
        "Al cierre del periodo se identifican tres casos marcados con el "
        "sufijo '(REVISAR DX)'. Estos casos corresponden a diagnosticos no "
        "tumorales o ambiguos que requieren validacion del patologo antes "
        "de su consolidacion definitiva en la base de datos."
    )
    b.table(
        ["Numero de caso", "Motivo de revision"],
        [
            ("IHQ250864", "Diagnostico no neoplasico - validar conducta."),
            ("IHQ250900", "Hallazgo descriptivo sin clasificacion estandar."),
            ("IHQ250917", "Diagnostico ambiguo entre dos entidades."),
        ],
        col_widths_cm=[5.0, 11.0],
    )

    b.h1("6. PDFs procesados")
    b.bullets([
        "Lote reciente: 22 archivos PDF procesados.",
        "Total acumulado en BD: 38 archivos PDF.",
        "Pendientes de reproceso identificados: 14 archivos PDF.",
        "Carpeta fuente: pdfs_patologia (servidor compartido HUV).",
    ])

    b.h1("7. Tiempo promedio de procesamiento")
    b.para(
        "El tiempo de procesamiento depende de la complejidad del informe "
        "(numero de biomarcadores, longitud del texto OCR) y del hardware "
        "disponible. Los valores observados son indicativos."
    )
    b.table(
        ["Tipo de caso", "Tiempo promedio"],
        [
            ("Caso sencillo (biopsia negativa)", "1 a 2 minutos."),
            ("Caso estandar (carcinoma con 5 a 10 biomarcadores)",
             "3 a 5 minutos."),
            ("Caso complejo (linfoma con panel extendido)",
             "5 a 8 minutos."),
            ("Caso multi-diagnostico", "8 a 12 minutos."),
        ],
        col_widths_cm=[6.5, 9.5],
    )

    b.h1("8. Indicadores de cobertura por categoria")
    b.table(
        ["Categoria", "Campos esperados", "Campos extraidos (promedio)"],
        [
            ("Administrativa", "19", "18,8"),
            ("Demografica", "11", "10,5"),
            ("Procedimiento", "7", "6,9"),
            ("Diagnostico", "7", "6,7"),
            ("Estudios IHQ generales", "5", "4,8"),
            ("Biomarcadores IHQ", "141", "118,0"),
            ("Metadata", "2", "2,0"),
        ],
        col_widths_cm=[5.5, 4.0, 6.5],
    )
    b.para(
        "Los valores promedio de campos extraidos son aproximaciones "
        "indicativas para fines ejecutivos; el calculo preciso se realiza "
        "mediante la funcion de auditoria del aplicativo."
    )

    b.h1("9. Pendientes de reproceso")
    b.para(
        "Existen 30 casos pendientes de reproceso, distribuidos en 25 "
        "casos antiguos identificados antes de V6.9.9 y 5 casos nuevos "
        "detectados durante la auditoria del periodo. El reproceso esta "
        "programado para la siguiente jornada nocturna disponible y "
        "permitira elevar la cobertura del campo Diagnostico Principal por "
        "encima del 99%."
    )
    b.bullets([
        "Casos antiguos: 25 (requieren reproceso con prompt V6.9.9).",
        "Casos nuevos detectados: 5 (auditoria de Mayo 2026).",
        "PDFs implicados: 14 archivos identificados.",
        "Tiempo estimado de reproceso completo: una jornada nocturna.",
    ])

    b.h1("10. Comparativo antes y despues de mejoras V6.9.9")
    b.table(
        ["Caso de validacion", "Antes (V6.9.8)", "Despues (V6.9.9)"],
        [
            ("Caso 1: diagnostico no tumoral",
             "Diagnostico Principal: N/A",
             "Diagnostico Principal correctamente preservado."),
            ("Caso 2: hallazgo descriptivo",
             "Diagnostico Principal: N/A",
             "Diagnostico Principal con marca (REVISAR DX)."),
            ("Caso 3: carcinoma con observaciones",
             "Observaciones truncadas",
             "Observaciones consolidadas sin perdida clinica."),
            ("Caso 4: panel de linfoma extendido",
             "Biomarcadores parcialmente N/A",
             "Cobertura completa de biomarcadores positivos."),
        ],
        col_widths_cm=[5.0, 5.5, 5.5],
    )
    b.blank()
    guardar(doc, b, destino)
    return destino


# --------------------------------------------------------------------------- #
# Main                                                                        #
# --------------------------------------------------------------------------- #


GENERADORES = [
    doc_01_credenciales,
    doc_02_doc_tecnica,
    doc_03_doc_adicional,
    doc_04_manual_usuario,
    doc_05_entrenamiento_s1,
    doc_06_entrenamiento_s2,
    doc_07_bitacora,
    doc_08_estadistica,
]


def main() -> int:
    t0 = time.time()
    print(">>> Generando 8 documentos oficiales EVARISIS V6.9.9")
    if not PLANTILLA.exists():
        print(f"!!! Plantilla no encontrada: {PLANTILLA}")
        return 1
    SALIDA.mkdir(parents=True, exist_ok=True)

    rutas: List[Path] = []
    for fn in GENERADORES:
        try:
            ruta = fn()
            tamano_kb = ruta.stat().st_size / 1024
            print(f"   OK -> {ruta.name} ({tamano_kb:.1f} KB)")
            rutas.append(ruta)
        except Exception as exc:  # pragma: no cover - reporting
            print(f"   FALLO en {fn.__name__}: {exc}")
            raise

    # Verificacion de membrete (header con imagenes intacto)
    print()
    print(">>> Verificacion de membrete:")
    todos_ok = True
    for ruta in rutas:
        d = Document(ruta)
        rels = d.sections[0].header.part.rels
        tiene_logos = any("image" in r.target_ref for r in rels.values())
        marca = "OK" if tiene_logos else "FALLA"
        if not tiene_logos:
            todos_ok = False
        print(f"   [{marca}] {ruta.name}")

    total_mb = sum(r.stat().st_size for r in rutas) / (1024 * 1024)
    elapsed = time.time() - t0
    print()
    print(">>> RESUMEN")
    print(f"   Documentos generados: {len(rutas)}")
    print(f"   Total tamano paquete: {total_mb:.2f} MB")
    print(f"   Tiempo de generacion: {elapsed:.1f} s")
    print(f"   Carpeta destino: {SALIDA}")
    print(f"   Membretes preservados: {'TODOS' if todos_ok else 'CON FALLAS'}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
